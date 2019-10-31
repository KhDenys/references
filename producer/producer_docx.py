import docx


def generate_docx(sources):
    document = docx.Document()

    document.add_heading('Список використаних джерел', 1)

    for source in sources:
        result = source_to_str(source)
        print(result)
        document.add_paragraph(
            result, style='List Number'
        )
    
    docx_file = 'sources.docx'
    document.save(docx_file)

    return docx_file


def source_to_str(source):
    book_name = source.name
    book_type = normalize_book_type(source.book_type)
    edition_city = normalize_book_city(source.city_of_publication)
    edition_number = f'{source.number}-е вид.'
    additional_info = normalize_book_info(source.extra_info)
    publisher = normalize_book_publisher(source.publisher)
    year = normalize_book_year(source.year)
    number_pages = normalize_book_pages(source.number_of_pages)

    authors = source.authors.all()

    authors_authors = [author for author in authors if author._type == 'au']

    editor_translator = editor_translator_to_str([author for author in authors if author._type != 'au'])


    
    if len(authors_authors) < 4:
        main_authors = []
        for author in authors_authors:
            main_authors.append(author_to_str(author))

        main_authors_to_str = (', '.join(main_authors) + ' ') if main_authors else ''
        return (
            f'{main_authors_to_str}'
            f'{book_name} '
            f'{book_type}'
            f'{editor_translator}'
            f'{edition_number}'
            f'{additional_info}'
            f'{edition_city}'
            f'{publisher}'
            f'{year}'
            f'{number_pages}'
        )

    else:
        main_author = author_to_str_reversed(authors_authors[0], many=True)

        return (
            f'{book_name} '
            f'{book_type}'
            f'{main_author}'
            f'{editor_translator}'
            f'{edition_number}'
            f'{additional_info}'
            f'{edition_city}'
            f'{publisher}'
            f'{year}'
            f'{number_pages}'
        )


def normalize_book_type(book_type):
    if book_type is None:
        return '/'

    book_type = book_type.strip().lower()
    book_type = book_type if book_type[-1] == '.' else (book_type + '.')

    return f': {book_type} /'


def normalize_book_city(city):
    if city is None:
        return ''
    
    return city.strip().capitalize()


def normalize_book_info(book_info):
    if book_info is None:
        return ''

    book_info = book_info.strip().lower()
    return f', {book_info} '


def editor_translator_to_str(editors_translators):
    if len(editors_translators) == 0:
        return ''

    author = editors_translators.pop()
    name = author_to_str_reversed(author)
    if author._type == 'ed':
        return f' за ред. {name} '
    else:
        return f' пер. з {author.lang} {name} '


def normalize_book_publisher(book_publisher):
    if book_publisher is None:
        return ''
    
    return f' : {book_publisher}'


def normalize_book_year(book_year):
    if book_year is None:
        return ','
    
    return f', {book_year}.'


def normalize_book_pages(number_pages):
    if number_pages is None:
        return ''
    
    return f' {number_pages} c.'


def author_to_str_reversed(author, many=False):
    name_reversed = author_to_str(author)
    name = ' '.join(reversed(name_reversed.split()))
    if many:
        name = f' {name} та ін.;'

    return name


def author_to_str(author):
    last_name = author.last_name.capitalize()
    first_name = author.first_name[0].upper()
    patronymic = author.patronymic[0].upper()
    return f'{last_name} {first_name}. {patronymic}.'
